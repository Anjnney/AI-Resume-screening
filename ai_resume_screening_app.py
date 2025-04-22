import os
import re
import nltk
import spacy
import streamlit as st
import sqlite3
import logging
from docx import Document
from pdfminer.high_level import extract_text
from pdf2image import convert_from_bytes
import pytesseract
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import string
nltk.download('punkt')
nltk.download('stopwords')
# 1. Define the function to recreate the table
def recreate_table():
    conn = sqlite3.connect('resumes.db')
    cursor = conn.cursor()

    # 2. Drop the table if it exists
    cursor.execute('DROP TABLE IF EXISTS resumes')

    # 3. Create the table with the correct schema
    cursor.execute('''
        CREATE TABLE resumes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT,
            email TEXT,
            phone TEXT,
            skills TEXT,
            experience TEXT,
            education TEXT,
            full_text TEXT
        )
    ''')

    conn.commit()
    conn.close()

# 4. Call the function to recreate the table (so it happens before any insert)
recreate_table()

# 5. Now, your insert operations and other logic go here
# For example, your insert function...

# Setup logging for duplicates and debug
logging.basicConfig(filename='duplicates.log', level=logging.INFO, format='%(asctime)s - %(message)s')
debug_log = logging.getLogger('debug')
debug_log.setLevel(logging.INFO)
fh = logging.FileHandler('scores.log')
fh.setLevel(logging.INFO)
debug_log.addHandler(fh)

# Download necessary NLTK resources
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('wordnet')

# Load spaCy model
nlp = spacy.load("en_core_web_sm")

# Skill dictionary for keyword enhancement
SKILLS_DICT = {
    # Programming Languages
    'python': 'Python', 'java': 'Java', 'c': 'C', 'c++': 'C++', 'c#': 'C#',
    'go': 'Go', 'kotlin': 'Kotlin', 'swift': 'Swift', 'typescript': 'TypeScript',
    'javascript': 'JavaScript', 'ruby': 'Ruby', 'r': 'R', 'rust': 'Rust',

    # Web Development - Frontend
    'html': 'HTML', 'css': 'CSS', 'react': 'React', 'angular': 'Angular',
    'vue': 'Vue.js', 'sass': 'SASS', 'bootstrap': 'Bootstrap',
    'tailwind': 'Tailwind CSS', 'webpack': 'Webpack', 'next.js': 'Next.js',
    'redux': 'Redux',

    # Web Development - Backend
    'node.js': 'Node.js', 'php': 'PHP', 'express': 'Express', 'django': 'Django',
    'flask': 'Flask', 'spring': 'Spring', 'laravel': 'Laravel', 'dotnet': '.NET',
    'rest': 'REST', 'graphql': 'GraphQL',

    # Databases
    'sql': 'SQL', 'mysql': 'MySQL', 'postgresql': 'PostgreSQL', 'mongodb': 'MongoDB',
    'redis': 'Redis', 'cassandra': 'Cassandra', 'hadoop': 'Hadoop',
    'spark': 'Apache Spark', 'hive': 'Hive', 'etl': 'ETL', 'data warehouse': 'Data Warehouse',

    # DevOps & Cloud Tools
    'docker': 'Docker', 'kubernetes': 'Kubernetes', 'ansible': 'Ansible',
    'terraform': 'Terraform', 'jenkins': 'Jenkins', 'github actions': 'GitHub Actions',
    'ci/cd': 'CI/CD', 'aws': 'AWS', 'azure': 'Azure', 'gcp': 'GCP',
    'cloudformation': 'CloudFormation', 'helm': 'Helm', 'prometheus': 'Prometheus',
    'grafana': 'Grafana',

    # Data Science & AI
    'machine learning': 'Machine Learning', 'deep learning': 'Deep Learning',
    'data science': 'Data Science', 'tensorflow': 'TensorFlow', 'keras': 'Keras',
    'pytorch': 'PyTorch', 'scikit-learn': 'Scikit-learn', 'xgboost': 'XGBoost',
    'nlp': 'NLP', 'openai': 'OpenAI', 'transformers': 'Transformers',
    'cv': 'Computer Vision', 'pandas': 'Pandas', 'numpy': 'NumPy',
    'matplotlib': 'Matplotlib', 'seaborn': 'Seaborn', 'huggingface': 'HuggingFace',
    'llm': 'LLM',

    # Cybersecurity
    'penetration testing': 'Penetration Testing', 'metasploit': 'Metasploit',
    'burpsuite': 'Burp Suite', 'nmap': 'Nmap', 'wireshark': 'Wireshark',
    'owasp': 'OWASP', 'kali linux': 'Kali Linux', 'snort': 'Snort',
    'siem': 'SIEM', 'threat hunting': 'Threat Hunting', 'incident response': 'Incident Response',
    'red team': 'Red Team', 'blue team': 'Blue Team', 'osint': 'OSINT',

    # Mobile Development
    'android': 'Android', 'ios': 'iOS', 'react native': 'React Native',
    'flutter': 'Flutter', 'xamarin': 'Xamarin', 'dart': 'Dart',

    # Testing & QA
    'selenium': 'Selenium', 'cypress': 'Cypress', 'junit': 'JUnit',
    'pytest': 'PyTest', 'testng': 'TestNG', 'postman': 'Postman',
    'loadrunner': 'LoadRunner', 'jmeter': 'JMeter',

    # Version Control & Agile
    'git': 'Git', 'github': 'GitHub', 'jira': 'JIRA',
    'agile': 'Agile', 'scrum': 'Scrum',

    # General Tools / BI / Analytics
    'excel': 'Excel', 'power bi': 'Power BI', 'tableau': 'Tableau',
    'google analytics': 'Google Analytics', 'seo': 'SEO', 'lookml': 'LookML',
    'qlikview': 'QlikView', 'data visualization': 'Data Visualization',
    'financial modeling': 'Financial Modeling'
}

# Custom stopwords (exclude job-relevant terms)
CUSTOM_STOPWORDS = set(stopwords.words('english')) - set(['python', 'java', 'sql', 'javascript', 'ai', 'machine', 'learning', 'data', 'science', 'excel', 'seo'])

# ---------------------- Resume Text Extraction ----------------------

def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text.append(cell.text)
    return '\n'.join([t for t in text if t.strip()])

def extract_text_from_pdf(uploaded_file):
    try:
        text = extract_text(uploaded_file)
        if text and len(text.strip()) > 50:
            return text
        uploaded_file.seek(0)
        images = convert_from_bytes(uploaded_file.read())
        text = ' '.join([pytesseract.image_to_string(img) for img in images])
        return text if text.strip() else "No text extracted from PDF."
    except Exception as e:
        return f"Error extracting text: {str(e)}"

# Function to validate if the resume is valid based on the extracted text
def is_valid_resume(text):
    if not text or len(text.strip()) < 100:  # Check if text is empty or too short
        return False
    # You can add more conditions if you want, like checking for certain keywords or sections
    return True


# ---------------------- NLP + Scoring ----------------------

def preprocess_text(text):
    tokens = [word.lower() for word in word_tokenize(text) if word.isalpha()]
    lemmatizer = WordNetLemmatizer()
    tokens = [lemmatizer.lemmatize(word) for word in tokens if word not in CUSTOM_STOPWORDS]
    skills = extract_skills(' '.join(tokens))
    cleaned_text = re.sub(r'\s+', ' ', ' '.join(tokens)).strip()
    return cleaned_text, skills

def extract_skills(text):
    found_skills = []
    text_lower = text.lower()
    for skill_key, skill_value in SKILLS_DICT.items():
        if skill_key in text_lower:
            found_skills.append(skill_value)
    return found_skills if found_skills else ['No skills found']

def extract_name(text):
    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            return ent.text
    return "Name not found"

def extract_email(text):
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text, re.IGNORECASE)
    return match.group(0) if match else "Email not found"

def extract_phone(text):
    match = re.search(r"(\+?\d{1,3}[-.\s]?)?(\(?\d{2,3}\)?[-.\s]?)?(\d{3}[-.\s]?)?\d{3}[-.\s]?\d{4}", text)
    return match.group(0) if match else "Phone number not found"

# ---------------------- Cosine Similarity Scoring ----------------------

def score_resume_with_job_requirements(resume_text, job_description):
    resume_cleaned_text, resume_skills = preprocess_text(resume_text)
    job_cleaned_text, job_skills = preprocess_text(job_description)
    
    # Vectorize texts
    tfidf_vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = tfidf_vectorizer.fit_transform([resume_cleaned_text, job_cleaned_text])
    cosine_sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
    similarity_score = cosine_sim[0][0]
    
    debug_log.info(f"Resume Skills: {', '.join(resume_skills)}")
    debug_log.info(f"Job Description Skills: {', '.join(job_skills)}")
    debug_log.info(f"Cosine Similarity Score: {similarity_score}")
    
    return similarity_score, resume_skills, job_skills
